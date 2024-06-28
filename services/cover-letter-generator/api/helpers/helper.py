from instance import config
from ..logger import logger
import re, os, json, random, spacy, requests
from datetime import datetime
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

embedding = spacy.load("en_core_web_md")
    
def get_available_titles_vectors():
    """
    Get available job titles vectors
    
    Returns:
        list: Available job titles vectors
    """
    titles_vectors = []
    for title in config.AVAILABLE_JOB_TITLES:
        title_vector = embedding(title).vector
        titles_vectors.append(title_vector)
        
    return titles_vectors
    
available_titles_embeddings = get_available_titles_vectors()

def fill_template_sentence(sentence, user_info, wanted_job_info, skills_counter, global_skills_counter, experiences_counter):
    """
    Fill a sentence template with user info and wanted job info

    Args:
        sentence (str): The sentence template
        user_info (dict): User info
        wanted_job_info (dict): Wanted job info
        skills_counter (int): Skills counter
        global_skills_counter (int): Global skills counter
        experiences_counter (int): Experiences counter
    Returns:
        str: The filled sentence
    """
    if re.search(r"\$worked_company", sentence) or re.search(r"\$worked_position", sentence) or re.search(r"\$company_exp_years", sentence):
        experiences_counter += 1
    
    # find all occurrences of variable $position, and replace them all with job title
    sentence = re.sub(r"\$position", wanted_job_info["jobTitle"], sentence)
    
    # find all occurrences of variable $applying_company, and replace them all with company name
    sentence = re.sub(r"\$applying_company", wanted_job_info["companyName"], sentence)
    
    if user_info.get("experiences") and len(user_info["experiences"]) > 0:
        # find all occurrences of variable $worked_company, and replace them all with user previous company name
        sentence = re.sub(r"\$worked_company", user_info["experiences"][0].get("companyName"), sentence)
        
        # find all occurrences of variable $worked_position, and replace them all with user previous job title
        sentence = re.sub(r"\$worked_position", user_info["experiences"][0].get("jobTitle"), sentence)
        
        # find all occurrences of variable $company_exp_years, and replace them all with user previous company experience years
        sentence = re.sub(r"\$company_exp_years", str(user_info["experiences"][0].get("companyExperienceYears")), sentence)
        
        # find all occurrences of variable $total_exp_years, and replace them all with user total experience years
        sentence = re.sub(r"\$total_exp_years", str(user_info["yearsOfExperience"]), sentence)
    
    
    # in a while loop, find all occurrences of variable $skill one by one, and replace them all with a skill from user skills list
    while re.search(r"\$skill", sentence):
        sentence = re.sub(r"\$skill", user_info["skills"][skills_counter], sentence, 1)
        
        # round the counter if it exceeds the length of the skills list
        skills_counter = (skills_counter + 1) % len(user_info["skills"])
        
        # just increment the global skills counter
        global_skills_counter += 1
        
    # add . at the end of the sentence if it doesn't exist
    if sentence[-1] != ".":
        sentence += "."
        
    return sentence, skills_counter, global_skills_counter, experiences_counter

def fill_experience_template(sentence, experience):
    """
    Fill a sentence template with experience info
    
    Args:
        sentence (str): The sentence template
        experience (dict): Experience info
    Returns:
        str: The filled sentence
    """
    # find all occurrences of variable $worked_company, and replace them all with user previous company name
    sentence = re.sub(r"\$worked_company", experience["companyName"], sentence)
    
    # find all occurrences of variable $worked_position, and replace them all with user previous job title
    sentence = re.sub(r"\$worked_position", experience["jobTitle"], sentence)
    
    # find all occurrences of variable $company_exp_years, and replace them all with user previous company experience years
    sentence = re.sub(r"\$company_exp_years", str(experience["companyExperienceYears"]), sentence)
    
    # add . at the end of the sentence if it doesn't exist
    if sentence[-1] != ".":
        sentence += "."
    
    return sentence

def fill_cover_letter(available_templates, user_info, wanted_job_info):
    """
    Fill a cover letter template with user info and wanted job info

    Args:
        available_templates (dict): Available templates
        user_info (dict): User info
        wanted_job_info (dict): Wanted job info
    Returns:
        str: The filled cover letter
    """
    
    # validate user_info
    if not user_info.get("skills") or len(user_info["skills"]) == 0:
        raise ValueError("User's skills list is empty")
    
    # Initialize the cover letter
    cover_letter = f"""{user_info["fullName"]}
{user_info["address"]}
{user_info["phoneNumber"]}
{user_info["email"]}
{datetime.now().strftime("%B %d, %Y")}
Dear Hiring Manager,
"""

    # skills counter, to cover all user's skills if the template has placeholders for all of them, and not only first some skills
    skills_counter = 0
    
    # global skills counter, doesn't wrap around, to keep track of the skills that were used in the cover letter
    global_skills_counter = 0
    
    # global experiences counter, doesn't wrap around, to keep track of the experiences that were used in the cover letter
    experiences_counter = 0
    
    sentences = []
    
    # Intro sentence
    # generate random number between 0 and len(available_templates["intro"]) - 1
    random_num = random.randint(0, len(available_templates["intro"]) - 1)
    intro_template = available_templates["intro"][random_num]
    intro_template, skills_counter, global_skills_counter, experiences_counter = fill_template_sentence(intro_template, user_info, wanted_job_info, skills_counter, global_skills_counter, experiences_counter)
    sentences.append(intro_template)
        
    # Experience sentence (remove if len(user_info["experiences"]) == 0)
    experience_template = None
    if user_info.get("experiences") and len(user_info["experiences"]) > 0:
        random_num = random.randint(0, len(available_templates["experience"]) - 1)
        experience_template = available_templates["experience"][random_num]
        experience_template, skills_counter, global_skills_counter, experiences_counter = fill_template_sentence(experience_template, user_info, wanted_job_info, skills_counter, global_skills_counter, experiences_counter)
        sentences.append(experience_template)
        
    # Skills sentence
    random_num = random.randint(0, len(available_templates["skills"]) - 1)
    skills_template = available_templates["skills"][random_num]
    skills_template, skills_counter, global_skills_counter, experiences_counter = fill_template_sentence(skills_template, user_info, wanted_job_info, skills_counter, global_skills_counter, experiences_counter)
    sentences.append(skills_template)
        
    # Closing sentence
    random_num = random.randint(0, len(available_templates["closing"]) - 1)
    closing_template = available_templates["closing"][random_num]
    closing_template, skills_counter, global_skills_counter, experiences_counter = fill_template_sentence(closing_template, user_info, wanted_job_info, skills_counter, global_skills_counter, experiences_counter)
    sentences.append(closing_template)
    
    # Insert skills sentence if user has more than skills_counter skills till 6
    if global_skills_counter < len(user_info["skills"]) - 1:
        random_num = random.randint(0, len(available_templates["additional_skills"]) - 1)
        skills_template = available_templates["additional_skills"][random_num]
        if experience_template is not None:
            skills_index = 3
        else:
            skills_index = 2
            
        while True:
            if global_skills_counter > len(user_info["skills"]) - 1 or global_skills_counter >= 6:
                break
            skills_template += user_info["skills"][global_skills_counter] + ", "
            global_skills_counter += 1
        
        # remove the last comma and space and replace them with a period
        skills_template = skills_template[:-2] + "."
        
        # replace last comma with ", and"
        parts = skills_template.rsplit(",", 1)
        skills_template = ", and".join(parts)
        
        # insert the skills sentence
        sentences.insert(skills_index, skills_template)
                     
    # Insert additional experience sentences if user has more than one previous job till 2
    if experience_template is not None:
        inserted_experience = False
        used_experience_template_num = None
        if experiences_counter == 0:
            random_num = random.randint(0, len(available_templates["additional_experiences"]) - 1)
            used_experience_template_num = random_num
            experience_template = available_templates["additional_experiences"][random_num]
            experience_template = fill_experience_template(experience_template, user_info["experiences"][0])
            experiences_counter += 1
            sentences.insert(2, experience_template)
            inserted_experience = True
            
        if len(user_info["experiences"]) > 1:
            random_num = random.randint(0, len(available_templates["additional_experiences"]) - 1)
            while used_experience_template_num == random_num:
                random_num = random.randint(0, len(available_templates["additional_experiences"]) - 1)

            experience_template = available_templates["additional_experiences"][random_num]
            experience_template = fill_experience_template(experience_template, user_info["experiences"][1])
            
            if inserted_experience:
                sentences.insert(3, experience_template)
            else:
                sentences.insert(2, experience_template)
    
    # Join all sentences
    cover_letter += "\n".join(sentences)
    
    # End the cover letter
    cover_letter += "\nSincerely,\n" + user_info["fullName"] + "."
        
    return cover_letter

def calculate_similarity(vec1, vec2):
    """
    Calculate cosine similarity between two vectors
    
    Args:
        vec1 (np.array): First vector
        vec2 (np.array): Second vector
    Returns:
        float: Cosine similarity
    """
    dot_product = np.dot(vec1, vec2)
    norm_vec1 = np.linalg.norm(vec1)
    norm_vec2 = np.linalg.norm(vec2)
    
    cosine_similarity = dot_product / (norm_vec1 * norm_vec2)
    return max(0, min(1, cosine_similarity))

def calculate_similarity_with_available_titles(wanted_job_title):
    """
    Calculate similarity between wanted job title and all available titles

    Args:
        wanted_job_title (str): Wanted job title
    Returns:
        list: Similarities between wanted job title and all available titles
    """
    logger.debug("wanted_job_title: %s", wanted_job_title)
    
    wanted_title_vector = embedding(wanted_job_title).vector
    
    similarities = []
    for vector in available_titles_embeddings:
        similarities.append(calculate_similarity(wanted_title_vector, vector))
        
    # sort similarities in descending order
    arg_sorted_similarities = np.argsort(similarities)[::-1]
    
    return similarities, arg_sorted_similarities

def get_similar_titles(wanted_job_title):
    """
    Get similar titles to the wanted job title
    
    Args:
        wanted_job_title (str): Wanted job title
    Returns:
        dict: Similar titles with similarity
    """
    wanted_job_title = wanted_job_title.lower()
    
    # calculate similarity between wanted job title and all available titles
    titles_similarities, arg_sorted_similarity = calculate_similarity_with_available_titles(wanted_job_title)
    
    # take the top 5 similar titles
    top_5_similar_titles = {config.INDEX_TO_AVAILABLE_TITLES[i].replace(" ", "-"): titles_similarities[i] for i in arg_sorted_similarity[:5]}

    # take from top_5_similar_titles all the titles with similarity > 0.85 (threshold)
    similar_titles_similarity = {title: similarity for title, similarity in top_5_similar_titles.items() if similarity > 0.85}
    
    # if no similar titles with similarity > 0.85, take the top 3 similar titles
    if len(similar_titles_similarity) == 0:
        similar_titles_similarity = {config.INDEX_TO_AVAILABLE_TITLES[i].replace(" ", "-"): titles_similarities[i] for i in arg_sorted_similarity[:3]}

    return similar_titles_similarity

def sample_titles(probabilities):
    """
    Sample titles based on probabilities
    
    Args:
        probabilities (dict): Titles probabilities
    Returns:
        dict: Sampled titles counts
    """
    sampled_counts = {}
    
    total_num_samples = 5 * len(probabilities)

    for title in probabilities.keys():
        # count occurrences of each title in sampled list
        sampled_counts[title] = int(probabilities[title] * total_num_samples)
        
    return sampled_counts

def write_to_word(cover_letter, filename):
    """
    Write the cover letter to a Word document
    
    Args:
        cover_letter (str): The cover letter
        filename (str): The filename
    """
    # Split the text into lines
    lines = cover_letter.split("\n")
    
    user_name = lines[0]
    
    # Create a new Word document
    doc = Document()

    # Add a title
    title = doc.add_heading(user_name, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Remove the user name from the lines
    lines = lines[1:]

    # adress, phone, email, & date
    for i in range(4):
        p = doc.add_paragraph(lines[i])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].italic = True
        
    p = doc.add_paragraph("Dear Hiring Manager,")
    p.paragraph_format.space_after = Pt(12)
    p.runs[0].bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    lines = lines[5:]
        
    # Add each line with appropriate formatting
    for i in range(len(lines) - 2):
        if lines[i]:
            p = doc.add_paragraph(lines[i])
            p.paragraph_format.space_after = Pt(12)
            
    # Add the last lines
    lines = lines[-2:]

    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.runs[0].bold = True

    # Save the document
    doc.save(filename)
    
def generate_cover_letter_data(user_info, wanted_job_info):
    """
    Generate a cover letter

    Args:
        user_info (dict): User info
        wanted_job_info (dict): Wanted job info
    Returns:
        None
    """
    similarities = get_similar_titles(wanted_job_info["jobTitle"])
    
    # normalize probabilities
    prob_sum = sum(list(similarities.values()))

    probabilities = {title: prob / prob_sum for title, prob in similarities.items()}

    logger.debug("Probabilities: %s", probabilities)

    titles_counts = sample_titles(probabilities)
        
    available_templates = {
        "intro": [],
        "experience": [],
        "skills": [],
        "closing": [],
    }
    titles_dir = os.path.join("dataset", "job-titles")
    # get all directories inside job-titles directory
    titles = os.listdir(titles_dir)
    for title in titles:
        title_dir = os.path.join(titles_dir, title)
        if not os.path.isdir(title_dir) or not title in titles_counts:
            continue
        with open(os.path.join(title_dir, f"{title}.json"), "r") as f:
            data = json.load(f)
            for key in data:
                available_templates[key].extend(data[key] * titles_counts[title])
    
    # shuffle all key values to get normal distribution of templates
    for key in available_templates:
        random.shuffle(available_templates[key])
    
    additional_dir = os.path.join("dataset", "additional")
    with open(os.path.join(additional_dir, "additional_skills.json"), "r") as f:
        additional_skills = json.load(f)
        available_templates["additional_skills"] = additional_skills["additional_skills"]
        
    with open(os.path.join(additional_dir, "additional_experiences.json"), "r") as f:
        additional_experiences = json.load(f)
        available_templates["additional_experiences"] = additional_experiences["additional_experiences"]
    
    try:
        filled_cover_letter = fill_cover_letter(available_templates, user_info, wanted_job_info)
    except Exception as e:
        logger.exception(e)
        return
    
    user_fullname = user_info["fullName"].replace(" ", "-")
    
    dir_path = "api/generated-coverletters/"

    # Create directory if it does not exist
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    
    filename = f'{dir_path}/{user_fullname}-{datetime.now().strftime("%d-%m-%Y,%H-%M-%S")}.docx'
        
    # print to docx
    write_to_word(filled_cover_letter, filename)
    
    return filled_cover_letter, filename

def preprocess_user_info(user_info):
    """
    Preprocess user info
    
    Args:
        user_info (dict): User info
    Returns:
        dict: Preprocessed user info
    """
    # for experiences, transform startDate and endDate to companyExperienceYears
    for experience in user_info["experiences"]:
        start_date = datetime.fromisoformat(experience["startDate"])
        
        if not experience.get("endDate"):
            end_date = datetime.now()
        else:
            end_date = datetime.fromisoformat(experience["endDate"])
            
        # calculate company experience years, using ceil
        experience["companyExperienceYears"] = (end_date - start_date).days // 365 + 1
        
    # sort by startDate, descendingly
    user_info["experiences"] = sorted(user_info["experiences"], key=lambda x: x["startDate"], reverse=True)
    
def upload_file(file_path):
    """
    Uploads the file to the uploader service
    
    Args:
        file_path (str): File path
    Returns:
        str: The uploaded file link
    """
    url = f"http://{config.SERVER_HOST}:3000/api/v1/uploader/upload"
    
    filename = file_path.split('/')[-1]

    with open(file_path, 'rb') as file:
        payload = {}
        files=[
            (
                'file',
                (
                    filename,
                    file
                )
            )
        ]
        
        headers = {}

        response = requests.request("POST", url, headers=headers, data=payload, files=files)

        logger.debug("upload response = %s" % response.text)
        
        return response.json()["link"]
 