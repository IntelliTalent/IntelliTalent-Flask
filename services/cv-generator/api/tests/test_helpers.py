import unittest
from unittest.mock import patch, MagicMock
from docx import Document
from api.helpers.helper import (
    get_month_year,
    heading,
    objective,
    education,
    experience,
    projects,
    skills,
    certificates,
    upload_file
)

class HelpersTest(unittest.TestCase):

    def test_get_month_year(self):
        self.assertEqual(get_month_year("Present"), "Present")
        self.assertEqual(get_month_year("2022-01-01"), "Jan 2022")
        
    def test_heading(self):
        document = Document()
        heading_content = {
            "fullName": "John Doe",
            "phoneNumber": "1234567890",
            "email": "john@example.com",
            "gitHub": "https://github.com/johndoe",
            "linkedIn": "https://linkedin.com/in/johndoe",
            "city": "City",
            "country": "Country"
        }
        heading(document, heading_content)
        self.assertEqual(len(document.paragraphs), 4)
        
        document = Document()
        heading_content = {
            "fullName": "John Doe",
            "phoneNumber": "1234567890",
            "email": "john@example.com",
            "linkedIn": "https://linkedin.com/in/johndoe",
            "city": "City",
            "country": "Country"
        }
        heading(document, heading_content)
        self.assertEqual(len(document.paragraphs), 4)
        
        document = Document()
        heading_content = {
            "fullName": "John Doe",
            "phoneNumber": "1234567890",
            "email": "john@example.com",
            "gitHub": "https://github.com/johndoe",
            "city": "City",
            "country": "Country"
        }
        heading(document, heading_content)
        self.assertEqual(len(document.paragraphs), 4)
        
    def test_objective(self):
        document = Document()
        objective_content = {
            "summary": "Seeking a challenging position in software development."
        }
        objective(document, objective_content)
        self.assertEqual(len(document.paragraphs), 2)

    def test_education(self):
        document = Document()
        education_content = [
            {
                "schoolName": "University of ABC",
                "degree": "Bachelor of Science in Computer Science",
                "startDate": "2020-09-01",
                "endDate": "2024-05-01",
                "description": "Studied various computer science subjects."
            }
        ]
        education(document, education_content)
        self.assertEqual(len(document.paragraphs), 4)

    def test_experience(self):
        document = Document()
        experience_content = [
            {
                "jobTitle": "Software Engineer",
                "companyName": "XYZ Corp",
                "startDate": "2022-01-01",
                "endDate": "Present",
                "description": "Developing software applications."
            }
        ]
        experience(document, experience_content)
        self.assertEqual(len(document.paragraphs), 4)

    def test_projects(self):
        document = Document()
        projects_content = [
            {
                "name": "Project A",
                "description": "Developed a web application using Django.",
                "skills": "Python, Django"
            }
        ]
        projects(document, projects_content)
        self.assertEqual(len(document.paragraphs), 3)

    def test_skills(self):
        document = Document()
        skills_content = [
            {
                "category": "Technologies",
                "list": "Python, Java, JavaScript"
            }
        ]
        skills(document, skills_content)
        self.assertEqual(len(document.paragraphs), 2)

    def test_certificates(self):
        document = Document()
        certificates_content = [
            {
                "title": "Certificate in Python Programming",
                "authority": "XYZ Institute",
                "issuedAt": "2023-03-15",
                "validUntil": "2024-03-15",
                "url": "https://example.com/certificate"
            }
        ]
        certificates(document, certificates_content)
        self.assertEqual(len(document.paragraphs), 4)
        
    @patch('api.helpers.helper.requests')
    def test_upload_file(self, requests_mock):
        file_path = "api/tests/test_files/test.docx"
        upload_url = "http://example.com/api/v1/uploader/upload"
        response_mock = MagicMock()
        response_mock.json.return_value = {"link": upload_url}
        requests_mock.request.return_value = response_mock

        uploaded_link = upload_file(file_path)

        requests_mock.request.assert_called_once()
        self.assertEqual(uploaded_link, upload_url)

if __name__ == '__main__':
    unittest.TextTestRunner(verbosity=2).run(unittest.TestLoader().loadTestsFromTestCase(HelpersTest))