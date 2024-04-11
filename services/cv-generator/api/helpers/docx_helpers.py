import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, Inches

# constants
HEADINGLINESPACE = 1.15
SUBHEADINGLINESPACE = 1.15
CONTENTLINESPACE = 1.5
SKILLWIDTH = 9

def insert_hr(paragraph):
    """
    Insert a horizontal line into a paragraph.
    
    Args:
        paragraph: The paragraph to insert the horizontal line into.
    """
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = docx.oxml.shared.OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = docx.oxml.shared.OxmlElement('w:bottom')
    bottom.set(docx.oxml.ns.qn('w:val'), 'single')
    bottom.set(docx.oxml.ns.qn('w:sz'), '6')
    bottom.set(docx.oxml.ns.qn('w:space'), '1')
    bottom.set(docx.oxml.ns.qn('w:color'), 'auto')
    pBdr.append(bottom)
    
def get_or_create_hyperlink_style(d):
    """
    Get or create the hyperlink style in a document.
    
    Args:
        d: The document to get or create the hyperlink style in.
    Returns:
        The name of the hyperlink style.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the hyperlink to.
        text: The text of the hyperlink.
        url: The URL of the hyperlink.
    Returns:
        The hyperlink element.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

def heading_style(paragraph, dual=True):
    """
    Apply the heading style to a paragraph.
    
    Args:
        paragraph: The paragraph to apply the heading style to.
        dual: Whether the heading is a dual heading.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = HEADINGLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(11)
    paragraph.runs[1].bold = False
    paragraph.runs[1].font.size = Pt(11)
    if dual:
        paragraph.runs[2].bold = True
        paragraph.runs[2].font.size = Pt(11)
        paragraph.runs[3].bold = False
        paragraph.runs[3].font.size = Pt(11)

def sub_heading_style(paragraph):
    """
    Apply the sub-heading style to a paragraph.
    
    Args:
        paragraph: The paragraph to apply the sub-heading style to.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = Cm(0.2)
    paragraph_format.line_spacing = SUBHEADINGLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(14)
    insert_hr(paragraph)

def content_heading_style(paragraph):
    """
    Apply the content heading style to a paragraph.
    
    Args:
        paragraph: The paragraph to apply the content heading style to.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = CONTENTLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(13)

def table_style(table):
    """
    Apply the table style to a table.
    
    Args:
        table: The table to apply the table style to.
    """
    table.autofit = False 
    table.allow_autofit = False
    table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table0_format = table.cell(0,0).paragraphs[0].paragraph_format
    table1_format = table.cell(0,1).paragraphs[0].paragraph_format
    table0_format.space_before = 0
    table0_format.space_after = 0
    table0_format.line_spacing = CONTENTLINESPACE
    table1_format.space_before = 0
    table1_format.space_after = 0
    table1_format.line_spacing = CONTENTLINESPACE
    table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in table.rows:
        row.height = Cm(0.1)

def bullet_list_style(paragraph):
    """
    Apply the bullet list style to a paragraph.
    
    Args:
        paragraph: The paragraph to apply the bullet list style to.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = CONTENTLINESPACE
    paragraph_format.left_indent = Inches(0.5)

def content_description_style(document, description):
    """
    Apply the content description style to a paragraph.
    
    Args:
        document: The document to apply the content description style to.
        description: The description text.
    """
    description = document.add_paragraph(description, style="List Bullet")
    bullet_list_style(description)
        

def skill_style(table):
    """
    Apply the skill style to a table.
    
    Args:
        table: The table to apply the skill style to.
    """
    table.autofit = False 
    table.allow_autofit = False
    for row in table.rows:
        row.height = Cm(0.1)
        row.cells[0].width = Cm(SKILLWIDTH)
        row.cells[1].width = 7052310 - SKILLWIDTH * 360000

        col0_format = row.cells[0].paragraphs[0].paragraph_format
        col1_format = row.cells[1].paragraphs[0].paragraph_format
        col0_format.space_before = 0
        col0_format.space_after = 0
        col0_format.line_spacing = CONTENTLINESPACE
        col1_format.space_before = 0
        col1_format.space_after = 0
        col1_format.line_spacing = CONTENTLINESPACE

def skill_heading_style(paragraph):
    """
    Apply the skill heading style to a paragraph.
    
    Args:
        paragraph: The paragraph to apply the skill heading style to.
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = CONTENTLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(11)
    