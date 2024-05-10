from .docx_helpers import *
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from instance import config
from ..logger import logger
import requests

def get_month_year(date):
    """
    Get the month and year from a date (Jan 2022).
    
    Args:
        date: The date to get the month and year from.
    Returns:
        str: The month and year.
    """
    # if it is str with val = Present or No expiry date, return it as is
    if date == "Present" or date == "No expiry date":
        return date
    
    date = datetime.fromisoformat(date)
    return date.strftime("%b %Y")

def heading(document, heading_content):
    """
    Add the heading to the document.
    
    Args:
        document: The document to add the heading to.
        heading_content: The content of the heading.
    """
    name = document.add_paragraph(heading_content["fullName"])
    name_format = name.paragraph_format
    name_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_format.space_before = name_format.space_after = 0
    name_format.line_spacing = 1
    name.runs[0].bold = True
    name.runs[0].font.size = Pt(26)
    
    second_line = document.add_paragraph("Phone: ")
    second_line.add_run(heading_content["phoneNumber"])
    second_line.add_run(" Email: ")
    second_line.add_run("")
    add_hyperlink(second_line, heading_content["email"], "mailto:" + heading_content["email"])
    heading_style(second_line)

    third_line = document.add_paragraph("Github: ")
    third_line.add_run("")
    add_hyperlink(third_line, heading_content["gitHub"], heading_content["gitHub"])
    third_line.add_run(" LinkedIn: ")
    third_line.add_run("")
    add_hyperlink(third_line, heading_content["linkedIn"], heading_content["linkedIn"])
    heading_style(third_line)
    
    fourth_line = document.add_paragraph("City: ")
    fourth_line.add_run(heading_content["city"])
    fourth_line.add_run(" Country: ")
    fourth_line.add_run(heading_content["country"])
    heading_style(fourth_line)
    
def objective(document, objective_content):
    """
    Add the objective to the document.
    
    Args:
        document: The document to add the objective to.
        objective_content: The content of the objective.
    """
    heading = document.add_paragraph("OBJECTIVE")
    objective = document.add_paragraph(objective_content["summary"] + "\n")
    objective_format = objective.paragraph_format
    objective_format.space_before = objective_format.space_after = 0
    objective_format.line_spacing = CONTENTLINESPACE
    objective.runs[0].bold = False
    objective.runs[0].font.size = Pt(11)
    sub_heading_style(heading)

def education(document, education_content):
    """
    Add the education to the document.
    
    Args:
        document: The document to add the education to.
        education_content: The content of the education.
    """
    heading = document.add_paragraph("EDUCATION")
    sub_heading_style(heading)
    for component in education_content:
        start_date = get_month_year(component["startDate"])
        end_date = get_month_year(component.get("endDate", "Present"))
        
        education = document.add_paragraph(component["schoolName"])
        table = document.add_table(rows=1, cols=2)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = component["degree"]
        heading_cells[1].text = start_date + " - " + end_date if start_date and end_date else ""
        table_style(table)

        content_description_style(document, component["description"])
        content_heading_style(education)
        document.add_paragraph("")
        
def experience(document, experience_content):
    """
    Add the experience to the document.
    
    Args:
        document: The document to add the experience to.
        experience_content: The content of the experience.
    """
    # subheading
    experience = document.add_paragraph("EXPERIENCE")
    sub_heading_style(experience)

    # add content
    for component in experience_content:
        start_date = get_month_year(component["startDate"])
        end_date = get_month_year(component.get("endDate", "Present"))
        
        experience = document.add_paragraph(component["jobTitle"])
        table = document.add_table(rows=1, cols=2)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = component["companyName"]
        if start_date and end_date:
            heading_cells[1].text = start_date + " - " + end_date
        else:
            heading_cells[1].text = ""
        table_style(table)
        content_heading_style(experience)
        content_description_style(document, component["description"])
        
        document.add_paragraph("")

def projects(document, projects_content):
    """
    Add the projects to the document.
    
    Args:
        document: The document to add the projects to.
        projects_content: The content of the projects.
    """
    projects_heading = document.add_paragraph("PROJECTS")
    sub_heading_style(projects_heading)
    for project in projects_content:
        project_table = document.add_table(rows=1, cols=2)
        project_cells = project_table.rows[0].cells
        project_cells[0].text = project["name"]
        content_heading_style(project_cells[0].paragraphs[0])
        table_style(project_table)
        content_description_style(document, project["description"])
        
        technologies_table = document.add_table(rows=0, cols=2)
        technologies_table.add_row()
        technologies_cells = technologies_table.rows[-1].cells
        technologies_cells[0].text = "Technologies"
        technologies_cells[1].text = project["skills"]
        skill_heading_style(technologies_cells[0].paragraphs[0])
        skill_style(technologies_table)
        document.add_paragraph("")

def skills(document, skills_content):
    """
    Add the skills to the document.
    
    Args:
        document: The document to add the skills to.
        skills_content: The content of the skills.
    """
    skills_heading = document.add_paragraph("SKILLS")
    sub_heading_style(skills_heading)
    skills_table = document.add_table(rows=0, cols=2)
    for skill in skills_content:
        skills_table.add_row()
        skill_cells = skills_table.rows[-1].cells
        skill_cells[0].text = skill["category"]
        skill_cells[1].text = skill["list"]
        skill_heading_style(skill_cells[0].paragraphs[0])
    skill_style(skills_table)
    document.add_paragraph("")
    
def certificates(document, certificates_content):
    """
    Add the certificates to the document.
    
    Args:
        document: The document to add the certificates to.
        certificates_content: The content of the certificates.
    """
    certificates_heading = document.add_paragraph("CERTIFICATES")
    sub_heading_style(certificates_heading)
    for certificate in certificates_content:
        issued_at = get_month_year(certificate["issuedAt"])
        valid_until = get_month_year(certificate.get("validUntil", "No expiry date"))
        
        certificate_paragraph = document.add_paragraph(certificate["title"])
        certificate_table = document.add_table(rows=1, cols=2)
        certificate_cells = certificate_table.rows[0].cells
        certificate_cells[0].text = certificate["authority"]
        certificate_cells[1].text = "Issued at: " + issued_at + " Valid until: " + valid_until
        table_style(certificate_table)
        content_heading_style(certificate_paragraph)
        
        certificate_url = document.add_paragraph("", style="List Bullet")
        add_hyperlink(certificate_url, certificate["url"], certificate["url"])
        bullet_list_style(certificate_url)
        document.add_paragraph("")
        
def upload_file(file_path):
    """
    Uploads the file to the uploader service
    
    Args:
        file_path (str): File path
    Returns:
        str: The uploaded file link
    """
    url = f"http://{config.SERVER_HOST}:3000/api/v1/uploader/upload"
    
    filename = file_path.split('/')[-1]

    payload = {}
    files=[
        (
            'file',
            (
                filename,
                open(file_path,'rb')
            )
        )
    ]
    
    headers = {}

    response = requests.request("POST", url, headers=headers, data=payload, files=files)

    logger.debug("upload response = %s" % response.text)
    
    return response.json()["link"]
 